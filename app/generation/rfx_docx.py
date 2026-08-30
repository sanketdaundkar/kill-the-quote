"""
Turns whatever RFx spec dict is currently loaded (the demo fixture or the
co-pilot's live draft - same shape either way, since the co-pilot's
update_rfx_draft tool schema mirrors rfx_spec.json) into a downloadable
Word document. This is the artifact a buyer would actually send to
vendors, so it's built to read like a real commercial document, not a
data dump.
"""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_background(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def build_rfx_docx_bytes(rfx: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    title = rfx.get("title") or "Request for Quote"
    category = rfx.get("category") or ""
    scope = rfx.get("scope") or ""
    currency = rfx.get("currency_base") or "INR"

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_bits = []
    if rfx.get("rfx_id"):
        meta_bits.append(f"Ref: {rfx['rfx_id']}")
    if category:
        meta_bits.append(f"Category: {category}")
    if rfx.get("issue_date"):
        meta_bits.append(f"Issued: {rfx['issue_date']}")
    if rfx.get("response_deadline"):
        meta_bits.append(f"Response due: {rfx['response_deadline']}")
    if meta_bits:
        p = doc.add_paragraph(" | ".join(meta_bits))
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    if scope:
        doc.add_heading("Scope", level=1)
        doc.add_paragraph(scope)

    terms = rfx.get("commercial_terms") or {}
    if any(terms.values()):
        doc.add_heading("Commercial Terms", level=1)
        term_labels = {
            "payment_terms": "Payment terms",
            "delivery_location": "Delivery location",
            "delivery_window": "Delivery window",
            "warranty_minimum": "Minimum warranty",
        }
        for key, label in term_labels.items():
            if terms.get(key):
                p = doc.add_paragraph(style=None)
                run = p.add_run(f"{label}: ")
                run.bold = True
                p.add_run(str(terms[key]))
        p = doc.add_paragraph()
        run = p.add_run(f"Currency: ")
        run.bold = True
        p.add_run(currency)

    line_items = rfx.get("line_items") or []
    if line_items:
        doc.add_heading("Line Items", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["Item Code", "Description", "Qty", "Unit"]
        for i, htext in enumerate(headers):
            hdr[i].text = htext
            hdr[i].paragraphs[0].runs[0].bold = True
            _set_cell_background(hdr[i], "2F5597")
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for idx, li in enumerate(line_items):
            row = table.add_row().cells
            row[0].text = str(li.get("item_code", ""))
            row[1].text = str(li.get("description", ""))
            row[2].text = str(li.get("qty", ""))
            row[3].text = str(li.get("unit", ""))
            if idx % 2 == 1:
                for c in row:
                    _set_cell_background(c, "F0F4FA")
        widths = [Inches(0.9), Inches(3.6), Inches(0.7), Inches(0.7)]
        for row in table.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = w

    questionnaire = rfx.get("questionnaire") or []
    if questionnaire:
        doc.add_heading("Vendor Questionnaire", level=1)
        for q in questionnaire:
            qid = q.get("q_id", "")
            qtext = q.get("question", "")
            doc.add_paragraph(f"{qid}: {qtext}" if qid else qtext, style="List Number")

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated with Aerchain's RFx co-pilot.")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
