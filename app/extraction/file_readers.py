"""
Reads vendor response files of any supported format and returns either:
  - plain text content (for xlsx/docx/pdf/txt/eml), or
  - a base64-encoded image payload (for jpg/png), for vision extraction.

This module deliberately does NOT try to interpret the content - it just
gets everything readable off the page/sheet/doc into a string (or image
bytes) so the LLM extraction step can do the actual understanding. Keeping
this dumb-and-honest is what lets extraction stay format-agnostic.
"""
import base64
import io
import os
from openpyxl import load_workbook
from docx import Document
import pdfplumber


def read_xlsx(path: str) -> str:
    wb = load_workbook(path, data_only=True)
    out = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        out.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                out.append(" | ".join(cells))
    return "\n".join(out)


def read_docx(path: str) -> str:
    doc = Document(path)
    out = []
    for para in doc.paragraphs:
        if para.text.strip():
            out.append(para.text)
    for table in doc.tables:
        out.append("--- Table ---")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            out.append(" | ".join(cells))
    return "\n".join(out)


def read_pdf(path: str) -> str:
    out = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            out.append(f"--- Page {i+1} ---")
            out.append(text)
            tables = page.extract_tables()
            for t in tables:
                out.append("--- Table ---")
                for row in t:
                    out.append(" | ".join([c or "" for c in row]))
    return "\n".join(out)


def read_txt(path: str) -> str:
    with open(path, "r", errors="ignore") as f:
        return f.read()


def read_image_b64(path: str) -> dict:
    with open(path, "rb") as f:
        data = f.read()
    ext = os.path.splitext(path)[1].lower().replace(".", "")
    media_type = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
    return {"media_type": media_type, "data": base64.standard_b64encode(data).decode("utf-8")}


def read_vendor_file(path: str):
    """Returns ('text', content) or ('image', {media_type, data})."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".xlsx":
        return "text", read_xlsx(path)
    if ext == ".docx":
        return "text", read_docx(path)
    if ext == ".pdf":
        return "text", read_pdf(path)
    if ext in (".txt", ".eml"):
        return "text", read_txt(path)
    if ext in (".jpg", ".jpeg", ".png"):
        return "image", read_image_b64(path)
    raise ValueError(f"Unsupported vendor file type: {ext}")
