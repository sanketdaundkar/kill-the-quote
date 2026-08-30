import json, random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(3)

with open('/home/claude/aerchain-rfx/data/rfx/rfx_spec.json') as f:
    rfx = json.load(f)
with open('/home/claude/aerchain-rfx/data/rfx/base_prices.json') as f:
    base = json.load(f)

desc_map = {li["item_code"]: li["description"] for li in rfx["line_items"]}
qty_map = {li["item_code"]: li["qty"] for li in rfx["line_items"]}

# Items Prime Hardware sells in boxes rather than per-unit (small accessories)
BOX_ITEMS = {
    "IT-007": 10,  # wired keyboard, box of 10
    "IT-008": 20,  # mouse, box of 20
    "IT-013": 25,  # HDMI cable, box of 25
    "IT-014": 50,  # Cat6 cable, box of 50
    "IT-023": 10,  # USB-C hub, box of 10
    "IT-024": 20,  # lock cable, box of 20
}

doc = Document()
h = doc.add_heading('Prime Hardware Co. — Commercial Offer', level=1)
p = doc.add_paragraph('Ref: PHC/QUOTE/2026-341   |   Date: 26-Aug-2026   |   In response to RFX-2026-0847')
p2 = doc.add_paragraph('Registered Office: Plot 14, MIDC Bhosari, Pune 411026, Maharashtra, GSTIN 27PRIMEH0123A1Z8')

doc.add_heading('1. Commercials', level=2)
doc.add_paragraph(
    "We are pleased to submit our best commercial terms below. Please note that for compute items "
    "(laptops, monitors, docking stations) we have quoted per-unit rates directly, exclusive of GST. "
    "For small-format accessories (keyboards, mice, cables, hubs and lock cables), owing to our packaging "
    "and logistics arrangement with our OEM partners, pricing is quoted per box as noted against each item; "
    "the buyer's team should compute per-unit equivalents based on box quantities mentioned. All other line "
    "items are quoted per unit. GST at 18% is additional throughout."
)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Item Code'
hdr[1].text = 'Description'
hdr[2].text = 'Qty Required'
hdr[3].text = 'Rate (INR, excl. GST)'

for code in desc_map:
    price = base[code] * random.uniform(0.92, 1.10)
    row = table.add_row().cells
    row[0].text = code
    row[1].text = desc_map[code]
    row[2].text = str(qty_map[code])
    if code in BOX_ITEMS:
        box_qty = BOX_ITEMS[code]
        box_price = round(price * box_qty, -1)
        row[3].text = f"Rs. {box_price:,.0f} per box of {box_qty}"
    else:
        row[3].text = f"Rs. {round(price, -1):,.0f} per unit"

doc.add_heading('2. Delivery & Warranty', level=2)
doc.add_paragraph(
    "Delivery within 5 weeks of PO confirmation, to Pune. Warranty is 3 years onsite for laptops and monitors, "
    "1 year OEM standard for all accessories and networking items. Freight to Pune is included in the above "
    "pricing for a single consolidated shipment; split shipments will attract additional freight of approx. "
    "Rs. 8,000-15,000 per shipment depending on volume."
)

doc.add_heading('3. Questionnaire Responses', level=2)
qa = doc.add_paragraph()
qa.add_run(
    "Q1 (OEM authorized reseller): Yes, for Dell, HP, Lenovo, TP-Link and APC.\n"
    "Q2 (Single Pune drop, one invoice): Yes.\n"
    "Q3 (On-site 3yr warranty support in Pune): Yes, via our authorized service partner network.\n"
    "Q4 (ISO 9001/27001 certified): ISO 9001 - yes. ISO 27001 - in progress, expected certification Q4 2026.\n"
    "Q5 (6-week delivery commitment): Yes, we are targeting 5 weeks."
)

doc.add_heading('4. Payment Terms', level=2)
doc.add_paragraph("Net 30 standard; Net 45 available for orders exceeding Rs. 75,00,000 subject to credit approval.")

doc.save('/home/claude/aerchain-rfx/data/vendor_responses/vendor_c_primehardware.docx')
print("saved")
